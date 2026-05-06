#!/usr/bin/env python3
"""
scripts/build_template.py

Genera la plantilla maestra SMNYL con estilos de marca aplicados y
placeholders Jinja2 listos para docxtpl.

Uso:
    python scripts/build_template.py

Salida:
    src/docs/templates/model_development_smnyl.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─── Tokens de marca SMNYL ────────────────────────────────────────────────────
_BLUE = "0079C2"    # NYL Blue — acentos, headers de tabla
_NAVY = "0A385E"    # Dark Rain — headers oscuros
_STEEL = "0A3C53"   # Texto principal
_IRON = "565656"    # Texto secundario / hints
_QUARTZ = "BDC1C2"  # Bordes
_WHITE = "FFFFFF"
_LIGHT = "F4F5F6"   # Fondo suave

FONT_D = "Georgia"  # display  (Alda Pro fallback)
FONT_B = "Tahoma"   # body     (Effra Pro fallback)

ROOT = Path(__file__).resolve().parent.parent
LOGO_PATH = ROOT / "assets" / "logo-smnyl.jpg"
OUT_PATH = ROOT / "src" / "docs" / "templates" / "model_development_smnyl.docx"


# ─── Primitivos de formato ────────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_bg(cell, hex6: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.lstrip("#"))
    tc_pr.append(shd)


def _table_borders(table, hex6: str = _QUARTZ) -> None:
    tbl_pr = table._tbl.tblPr
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex6.lstrip("#"))
        bdr.append(el)
    tbl_pr.append(bdr)


def _spacing(para, before: int = 0, after: int = 6) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_run(para, text: str, font: str = FONT_B, size: int = 11,
             color: str = _STEEL, bold: bool = False, italic: bool = False) -> None:
    r = para.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    r.font.bold = bold
    r.font.italic = italic


def _left_bar(para, color: str = _BLUE, width_pts: int = 24) -> None:
    """Barra vertical izquierda sobre un párrafo."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pts))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)


def _hr(doc: Document, color: str = _BLUE) -> None:
    p = doc.add_paragraph()
    _spacing(p, 4, 4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── Bloques tipográficos ─────────────────────────────────────────────────────

def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    r.font.name = FONT_D
    r.font.size = Pt(18)
    r.font.color.rgb = _rgb(_STEEL)
    r.font.bold = True
    _spacing(p, before=16, after=6)
    p.paragraph_format.left_indent = Cm(0.55)
    _left_bar(p, _BLUE, 24)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    r.font.name = FONT_D
    r.font.size = Pt(13)
    r.font.color.rgb = _rgb(_STEEL)
    r.font.bold = True
    _spacing(p, before=10, after=4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    r.font.name = FONT_B
    r.font.size = Pt(11)
    r.font.color.rgb = _rgb(_STEEL)
    r.font.bold = True
    _spacing(p, before=8, after=3)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _add_run(p, text)
    _spacing(p, 0, 6)


def _hint(doc: Document, text: str) -> None:
    """Instrucción en cursiva para el editor — no es contenido final."""
    p = doc.add_paragraph()
    _add_run(p, text, color=_IRON, italic=True)
    _spacing(p, 0, 4)


def _ph(doc: Document, var: str) -> None:
    """Bloque de texto con placeholder Jinja2 (azul info para visibilidad)."""
    p = doc.add_paragraph()
    _add_run(p, "{{ " + var + " }}", color="2E86AF")  # Medium Rain
    _spacing(p, 0, 10)


# ─── Setup de documento ───────────────────────────────────────────────────────

def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.0)
    sec.different_first_page_header_footer = True


def _setup_header_footer(doc: Document) -> None:
    sec = doc.sections[0]

    # Header (páginas 2+)
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    if LOGO_PATH.exists():
        lr = hp.add_run()
        lr.add_picture(str(LOGO_PATH), height=Cm(0.75))
    hp.add_run("\t")
    tr = hp.add_run("Model Development Documentation")
    tr.font.name = FONT_B
    tr.font.size = Pt(9)
    tr.font.color.rgb = _rgb(_IRON)
    # Línea bajo el header
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _QUARTZ)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer (páginas 2+)
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    # Línea sobre el footer
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"), "single")
    top_b.set(qn("w:sz"), "4")
    top_b.set(qn("w:space"), "1")
    top_b.set(qn("w:color"), _QUARTZ)
    fpBdr.append(top_b)
    fpPr.append(fpBdr)

    dr = fp.add_run("Borrador asistido por IA — Requiere revisión humana")
    dr.font.name = FONT_B
    dr.font.size = Pt(8)
    dr.font.color.rgb = _rgb(_IRON)
    dr.font.italic = True
    fp.add_run("\t")

    # Campo PAGE
    pr = fp.add_run("Pág. ")
    pr.font.name = FONT_B
    pr.font.size = Pt(8)
    pr.font.color.rgb = _rgb(_IRON)
    for tag, instr in (("begin", ""), ("", " PAGE "), ("end", "")):
        if tag:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), tag)
            pr._r.append(fld)
        else:
            it = OxmlElement("w:instrText")
            it.text = instr
            pr._r.append(it)

    # Footer portada (logo + empresa)
    fftr = sec.first_page_footer
    fftr.is_linked_to_previous = False
    ffp = fftr.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        ffp.add_run().add_picture(str(LOGO_PATH), height=Cm(1.0))
    cr = ffp.add_run("\nSeguros Monterrey New York Life  |  Confidencial")
    cr.font.name = FONT_B
    cr.font.size = Pt(8)
    cr.font.color.rgb = _rgb(_IRON)


# ─── Portada ──────────────────────────────────────────────────────────────────

def _add_cover(doc: Document) -> None:
    for _ in range(4):
        _spacing(doc.add_paragraph(), 0, 0)

    p_tipo = doc.add_paragraph()
    _add_run(p_tipo, "MODEL DEVELOPMENT TEMPLATE", FONT_B, 10, _IRON)
    _spacing(p_tipo, 0, 4)

    p_modelo = doc.add_paragraph()
    r_m = p_modelo.add_run("{{ nombre_modelo }}")
    r_m.font.name = FONT_D
    r_m.font.size = Pt(28)
    r_m.font.color.rgb = _rgb(_STEEL)
    _spacing(p_modelo, 4, 4)

    _hr(doc, _BLUE)

    for label, ph in (
        ("Preparado por", "{{ autor }}"),
        ("Fecha", "{{ fecha_documentacion }}"),
        ("Versión", "{{ version_actual }}"),
        ("Estado", "{{ estado_documento }}"),
    ):
        p = doc.add_paragraph()
        _add_run(p, f"{label}:   ", FONT_B, 10, _IRON, bold=True)
        _add_run(p, ph, FONT_B, 10, _STEEL)
        _spacing(p, 0, 3)

    for _ in range(7):
        _spacing(doc.add_paragraph(), 0, 0)

    doc.add_page_break()


# ─── Sección 1: Model Profile ─────────────────────────────────────────────────

def _add_section_1(doc: Document) -> None:
    _h1(doc, "1. Model Profile")

    # 1.1 Tabla de atributos
    _h2(doc, "1.1  Model Attributes")
    _hint(doc, "Metadata estructurada del modelo. Completar todos los campos antes de pasar a revisión.")

    attrs = [
        ("Model Name", "{{ model_name }}"),
        ("Model ID", "{{ model_id }}"),
        ("Model Class", "{{ model_class }}"),
        ("Profit Center", "{{ profit_center }}"),
        ("BU Executive (FAE)", "{{ bu_executive }}"),
        ("Model Owner", "{{ model_owner }}"),
        ("Model Developer(s)", "{{ model_developers }}"),
        ("Current Model Version", "{{ current_version }}"),
        ("Financial Impact", "{{ financial_impact }}"),
        ("Model Status", "{{ model_status }}"),
        ("Target Production Date", "{{ target_production_date }}"),
    ]
    tbl = doc.add_table(rows=len(attrs), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, ph) in enumerate(attrs):
        r = tbl.rows[i]
        cl = r.cells[0]
        cl.width = Cm(6.5)
        _cell_bg(cl, _NAVY)
        rl = cl.paragraphs[0].add_run(label)
        rl.font.name = FONT_B
        rl.font.size = Pt(10)
        rl.font.color.rgb = _rgb(_WHITE)
        rl.font.bold = True
        cv = r.cells[1]
        rv = cv.paragraphs[0].add_run(ph)
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    _table_borders(tbl)
    doc.add_paragraph()

    # 1.2 Version Control
    _h2(doc, "1.2  Version Control")
    _hint(doc, "Bitácora de cambios. DocuMente la puebla automáticamente desde el audit trail.")
    vc_headers = ["Version No.", "Date Changed", "Updated By", "Approved By", "Description"]
    vc = doc.add_table(rows=4, cols=5)
    vc.style = "Table Grid"
    # Header
    for i, h in enumerate(vc_headers):
        _cell_bg(vc.rows[0].cells[i], _BLUE)
        r = vc.rows[0].cells[i].paragraphs[0].add_run(h)
        r.font.name = FONT_B
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(_WHITE)
        r.font.bold = True
    # Loop start
    for c in vc.rows[1].cells:
        _cell_bg(c, "F0F0F0")
    lr = vc.rows[1].cells[0].paragraphs[0].add_run("{%- tr for v in version_history %}")
    lr.font.size = Pt(8)
    lr.font.color.rgb = _rgb("AAAAAA")
    # Data row
    for i, f in enumerate(["v.version_no", "v.date_changed", "v.updated_by", "v.approved_by", "v.description"]):
        rv = vc.rows[2].cells[i].paragraphs[0].add_run("{{ " + f + " }}")
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    # Loop end
    for c in vc.rows[3].cells:
        _cell_bg(c, "F0F0F0")
    er = vc.rows[3].cells[0].paragraphs[0].add_run("{%- tr endfor %}")
    er.font.size = Pt(8)
    er.font.color.rgb = _rgb("AAAAAA")
    _table_borders(vc)
    doc.add_paragraph()

    # 1.3 Problem Statement
    _h2(doc, "1.3  Problem Statement")
    _hint(doc, "Descripción del problema que el modelo resuelve. Incluir: limitaciones del modelo anterior, constraints de plataforma y productos cubiertos.")
    _ph(doc, "problem_statement")


# ─── Sección 2: Model Overview ────────────────────────────────────────────────

def _add_section_2(doc: Document) -> None:
    _h1(doc, "2. Model Overview")

    _h2(doc, "2.1  Model Uses")
    _hint(doc, "Descripción de todos los usuarios intencionados, frecuencia de uso y usos explícitamente fuera del alcance.")
    _ph(doc, "model_uses")

    _h2(doc, "2.2  Model Scope")
    _hint(doc, "Productos modelados: descripción de alto nivel y detalle de generaciones. Para cada producto: features que se modelan y features que se omiten (con justificación).")
    _ph(doc, "model_scope")

    _h2(doc, "2.3  Business Impact of Model Usage")
    _hint(doc, "Cómo encaja el modelo en las decisiones del negocio. Indicar si los resultados son requeridos por regulación y qué decisiones dependen del output.")
    _ph(doc, "business_impact")


# ─── Sección 3: Related & Supporting Documents ───────────────────────────────

def _add_section_3(doc: Document) -> None:
    _h1(doc, "3. Related & Supporting Documents")

    _h2(doc, "3.1  Ancillary Document List")
    _hint(doc, "Lista de documentos relacionados: carpeta del proyecto, ubicación de specs del modelo, sub-folders relevantes.")
    _ph(doc, "ancillary_docs")

    _h2(doc, "3.2  Additional Documents")
    _hint(doc, "Documentación adicional relacionada: pricing memos, policy forms, valuation memos, assumption memos, validation memos, etc.")
    _ph(doc, "additional_docs")


# ─── Sección 4: Model Description & Concept ──────────────────────────────────

def _add_section_4(doc: Document) -> None:
    _h1(doc, "4. Model Description & Concept")

    _h2(doc, "4.1  Schematic Diagram")
    _hint(doc, "Diagrama del sistema: fuentes de datos, modelos upstream, outputs, plataforma. Insertar imagen o describir verbalmente.")
    # Placeholder de imagen
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = p_img.add_run("[ Insertar diagrama esquemático del modelo ]")
    r_img.font.name = FONT_B
    r_img.font.size = Pt(10)
    r_img.font.color.rgb = _rgb(_QUARTZ)
    r_img.font.italic = True
    _spacing(p_img, 8, 8)
    _ph(doc, "schematic_description")

    _h2(doc, "4.2  Model Theory and Logic")
    _hint(doc, "Algoritmo(s) central(es) del modelo. Si usa plataforma existente: lógica adicional desarrollada. Enfoques alternativos considerados y razón por la que no se eligieron.")
    _ph(doc, "theory_and_logic")

    _h2(doc, "4.3  Key Risk Drivers")
    _hint(doc, "Listado y contexto de los 3-5 drivers principales de riesgo. Para cada uno: por qué es relevante y cómo está modelado.")
    _ph(doc, "key_risk_drivers")

    _h2(doc, "4.4  Key Assumptions")
    _hint(doc, "Todos los supuestos del modelo (económicos, actuariales). Para cada uno: fuente que lo respalda, rango plausible, y si podría necesitar revisión.")
    _ph(doc, "key_assumptions")


# ─── Sección 5: Inputs and Data ──────────────────────────────────────────────

def _add_section_5(doc: Document) -> None:
    _h1(doc, "5. Inputs and Data")

    # 5.1 Raw Data Sources
    _h2(doc, "5.1  Raw Data Sources and Data Quality")
    _hint(doc, "Para cada fuente de datos: tipo, proveedor, equipo responsable, ubicación, método de ingesta y análisis de calidad (accuracy, completeness, conformity).")

    raw_headers = ["Data Type", "Data Source", "Team Responsible", "Location", "Method"]
    raw = doc.add_table(rows=4, cols=5)
    raw.style = "Table Grid"
    for i, h in enumerate(raw_headers):
        _cell_bg(raw.rows[0].cells[i], _BLUE)
        r = raw.rows[0].cells[i].paragraphs[0].add_run(h)
        r.font.name = FONT_B
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(_WHITE)
        r.font.bold = True
    for c in raw.rows[1].cells:
        _cell_bg(c, "F0F0F0")
    raw.rows[1].cells[0].paragraphs[0].add_run("{%- tr for d in raw_data_sources %}").font.size = Pt(8)
    for i, f in enumerate(["d.data_type", "d.data_source", "d.team", "d.location", "d.method"]):
        rv = raw.rows[2].cells[i].paragraphs[0].add_run("{{ " + f + " }}")
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    for c in raw.rows[3].cells:
        _cell_bg(c, "F0F0F0")
    raw.rows[3].cells[0].paragraphs[0].add_run("{%- tr endfor %}").font.size = Pt(8)
    _table_borders(raw)
    doc.add_paragraph()

    # 5.2 Upstream Models
    _h2(doc, "5.2  Upstream Models & Company Determined Assumptions")
    _hint(doc, "Lista de todos los modelos upstream y supuestos determinados por la compañía.")

    up_headers = ["#", "Upstream Model / Assumption", "Key Contact", "Inventory ID"]
    up = doc.add_table(rows=4, cols=4)
    up.style = "Table Grid"
    for i, h in enumerate(up_headers):
        _cell_bg(up.rows[0].cells[i], _BLUE)
        r = up.rows[0].cells[i].paragraphs[0].add_run(h)
        r.font.name = FONT_B
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(_WHITE)
        r.font.bold = True
    for c in up.rows[1].cells:
        _cell_bg(c, "F0F0F0")
    up.rows[1].cells[0].paragraphs[0].add_run("{%- tr for u in upstream_models %}").font.size = Pt(8)
    for i, f in enumerate(["u.num", "u.name", "u.key_contact", "u.inventory_id"]):
        rv = up.rows[2].cells[i].paragraphs[0].add_run("{{ " + f + " }}")
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    for c in up.rows[3].cells:
        _cell_bg(c, "F0F0F0")
    up.rows[3].cells[0].paragraphs[0].add_run("{%- tr endfor %}").font.size = Pt(8)
    _table_borders(up)
    doc.add_paragraph()

    # 5.3 Pre-Processing
    _h2(doc, "5.3  Key Data Pre-Processing Steps")
    _h3(doc, "5.3.1  Data Aggregations")
    _hint(doc, "Bloques de pólizas tratados como unidad homogénea. Bloques con features similares tratados como idénticos.")
    _ph(doc, "data_aggregations")

    _h3(doc, "5.3.2  Segmentations")
    _hint(doc, "Segmentaciones relevantes: conversión de tasas, bloques por producto, segmentaciones regulatorias.")
    _ph(doc, "segmentations")

    _h3(doc, "5.3.3  Use of Averages or Proxies")
    _hint(doc, "Dónde se usan promedios o proxies para llenar datos faltantes o suavizar outliers.")
    _ph(doc, "averages_proxies")

    # 5.4 Known Limitations
    _h2(doc, "5.4  Known Input and Data Limitations")
    _hint(doc, "Lista de todas las limitaciones encontradas en datos o supuestos y las acciones tomadas para remediarlas.")
    _ph(doc, "data_limitations")

    # 5.5 Record of Input Changes
    _h2(doc, "5.5  Record of Input Changes or Decisions Made")
    _hint(doc, "Bitácora viva de cambios en inputs y decisiones. DocuMente la puebla desde el audit trail.")
    ic_headers = ["Date", "Decision / Change"]
    ic = doc.add_table(rows=4, cols=2)
    ic.style = "Table Grid"
    for i, h in enumerate(ic_headers):
        _cell_bg(ic.rows[0].cells[i], _BLUE)
        r = ic.rows[0].cells[i].paragraphs[0].add_run(h)
        r.font.name = FONT_B
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(_WHITE)
        r.font.bold = True
    for c in ic.rows[1].cells:
        _cell_bg(c, "F0F0F0")
    ic.rows[1].cells[0].paragraphs[0].add_run("{%- tr for c in input_changes %}").font.size = Pt(8)
    for i, f in enumerate(["c.date", "c.decision"]):
        rv = ic.rows[2].cells[i].paragraphs[0].add_run("{{ " + f + " }}")
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    for c in ic.rows[3].cells:
        _cell_bg(c, "F0F0F0")
    ic.rows[3].cells[0].paragraphs[0].add_run("{%- tr endfor %}").font.size = Pt(8)
    _table_borders(ic)
    doc.add_paragraph()


# ─── Sección 6: Model Build Process ──────────────────────────────────────────

def _add_section_6(doc: Document) -> None:
    _h1(doc, "6. Model Build Process")

    _h2(doc, "6.1  Specification Process")
    _hint(doc, "Si es nuevo: decisiones de negocio y constraints que condicionan las specs. Si es cambio: naturaleza del cambio respecto al modelo anterior, qué es nuevo.")
    _ph(doc, "specification_process")

    _h2(doc, "6.2  Approach Used")
    _hint(doc, "Metodología detallada: teoría y lógica, variable selection, derivaciones analíticas, supuestos implícitos en la estructura, comparación con enfoques alternativos (pros/cons) y racional de la elección final.")
    _ph(doc, "approach_used")

    _h2(doc, "6.3  Development Testing")
    _hint(doc, "Tests realizados: accuracy, robustness, sensitivity analysis, scenario testing (incluyendo extremos). Si hay test plan formal, documentar resultados de cada test.")
    _ph(doc, "development_testing")

    _h2(doc, "6.4  Limitations Revealed During Testing")
    _hint(doc, "Limitaciones descubiertas: por naturaleza de supuestos, simplificaciones implícitas o explícitas, procesos involucrados en producir resultados.")
    _ph(doc, "limitations_revealed")

    _h2(doc, "6.5  Record of Process Changes")
    _hint(doc, "Bitácora viva de cambios de proceso durante el build. DocuMente la puebla desde el audit trail.")
    pc_headers = ["Date", "Change Description"]
    pc = doc.add_table(rows=4, cols=2)
    pc.style = "Table Grid"
    for i, h in enumerate(pc_headers):
        _cell_bg(pc.rows[0].cells[i], _BLUE)
        r = pc.rows[0].cells[i].paragraphs[0].add_run(h)
        r.font.name = FONT_B
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(_WHITE)
        r.font.bold = True
    for c in pc.rows[1].cells:
        _cell_bg(c, "F0F0F0")
    pc.rows[1].cells[0].paragraphs[0].add_run("{%- tr for c in process_changes %}").font.size = Pt(8)
    for i, f in enumerate(["c.date", "c.change_description"]):
        rv = pc.rows[2].cells[i].paragraphs[0].add_run("{{ " + f + " }}")
        rv.font.name = FONT_B
        rv.font.size = Pt(10)
        rv.font.color.rgb = _rgb(_STEEL)
    for c in pc.rows[3].cells:
        _cell_bg(c, "F0F0F0")
    pc.rows[3].cells[0].paragraphs[0].add_run("{%- tr endfor %}").font.size = Pt(8)
    _table_borders(pc)
    doc.add_paragraph()


# ─── Sección 7: Model Implementation & Production ────────────────────────────

def _add_section_7(doc: Document) -> None:
    _h1(doc, "7. Model Implementation & Production")

    _h2(doc, "7.1  Platform")
    _hint(doc, "Plataforma donde corre el modelo (Prophet, GGY Axis, R, Python, Excel). Fuentes de datos en producción, proceso de transferencia, outputs y dónde se almacenan.")
    _ph(doc, "platform")

    _h2(doc, "7.2  Model Runs")
    _hint(doc, "Instrucciones de ejecución por caso de uso, settings de control, periodicidad de runs por propósito.")
    _ph(doc, "model_runs")

    _h2(doc, "7.3  Performance Testing")
    _hint(doc, "Testing para asegurar que el modelo refleja correctamente las specs y se desempeña como se espera bajo rangos relevantes.")
    _ph(doc, "performance_testing")

    _h2(doc, "7.4  Production and Performance Limitations")
    _hint(doc, "Limitaciones en producción: bajo qué condiciones el modelo podría desempeñarse inadecuadamente, aspectos del ambiente productivo cuyo cambio podría romper el modelo.")
    _ph(doc, "production_limitations")


# ─── Sección 8: Model Governance ─────────────────────────────────────────────

def _add_section_8(doc: Document) -> None:
    _h1(doc, "8. Model Governance")
    _hint(doc, "Controles sobre software, datos y outputs: version control, access control, IT controls. Controles sobre input data (protección contra tampering), transmisión de outputs, signoff authority y procedimientos de escalación.")
    _ph(doc, "governance")


# ─── Sección 9: On-going Monitoring ──────────────────────────────────────────

def _add_section_9(doc: Document) -> None:
    _h1(doc, "9. On-going Monitoring")
    _hint(doc, "Procedimientos de monitoreo de performance: KPIs del modelo, quién monitorea, frecuencia, thresholds y acciones cuando un threshold se rompe.")
    _ph(doc, "ongoing_monitoring")


# ─── Entry point ─────────────────────────────────────────────────────────────

def build() -> None:
    doc = Document()

    _setup_page(doc)
    _setup_header_footer(doc)
    _add_cover(doc)

    _add_section_1(doc)
    _add_section_2(doc)
    _add_section_3(doc)
    _add_section_4(doc)
    _add_section_5(doc)
    _add_section_6(doc)
    _add_section_7(doc)
    _add_section_8(doc)
    _add_section_9(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK  Plantilla generada: {OUT_PATH}")


if __name__ == "__main__":
    build()

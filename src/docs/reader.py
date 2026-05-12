"""DocxReader: parsea un .docx y lo convierte en un `Documento` estructurado.

Estrategia:
1. Leer todos los párrafos del .docx en orden.
2. Detectar headings que coincidan (por número o por nombre/aliases) con
   las secciones del catálogo oficial NYL (`template_catalog.py`).
3. Acumular el contenido entre cada heading detectado y la siguiente sección
   detectada como cuerpo de esa sección.
4. Producir un `Documento` con todas las secciones del catálogo: las que no
   se detectaron quedan vacías; las que sí, con su contenido.

Esto NO depende de que el .docx use el template oficial al pie de la letra
— los aliases en el catálogo permiten reconocer variantes razonables.
"""

from __future__ import annotations

import re
from datetime import UTC, datetime
from pathlib import Path

from docx import Document as DocxDocument

from src.core.models import Documento, EventoAuditoria, Seccion
from src.core.template_catalog import (
    TEMPLATE_MODEL_DEVELOPMENT,
    SeccionCatalogo,
    construir_secciones_vacias,
)

_NUM_PREFIX_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+(.+)$")


def _normalizar(texto: str) -> str:
    """Lowercase + strip + colapsa espacios para comparación de aliases."""
    return re.sub(r"\s+", " ", texto.strip().lower())


def _es_heading(estilo: str | None) -> bool:
    """True si el estilo de párrafo es un heading."""
    if not estilo:
        return False
    return estilo.lower().startswith("heading") or estilo.lower().startswith("title")


def _coincide_con_catalogo(texto_heading: str) -> SeccionCatalogo | None:
    """Intenta mapear un heading del .docx a una sección del catálogo.

    Estrategia (en orden de preferencia):
    1. Coincidencia por número exacto al inicio del heading (ej. "4.4 ...").
    2. Coincidencia exacta del texto normalizado contra nombre o alias.
    3. Coincidencia por substring del nombre normalizado.
    """
    texto_norm = _normalizar(texto_heading)
    if not texto_norm:
        return None

    # 1. Número al inicio
    m = _NUM_PREFIX_RE.match(texto_heading)
    if m:
        numero = m.group(1)
        for sec in TEMPLATE_MODEL_DEVELOPMENT:
            if sec.numero == numero:
                return sec

    # 2. Match exacto contra nombre o alias
    for sec in TEMPLATE_MODEL_DEVELOPMENT:
        if texto_norm == _normalizar(sec.nombre):
            return sec
        for alias in sec.aliases:
            if texto_norm == _normalizar(alias):
                return sec

    # 3. Substring (más permisivo)
    for sec in TEMPLATE_MODEL_DEVELOPMENT:
        nombre_norm = _normalizar(sec.nombre)
        if nombre_norm and nombre_norm in texto_norm:
            return sec

    return None


def _evaluar_completitud(contenido: str) -> str:
    """Heurística simple para clasificar completitud por longitud."""
    n_chars = len(contenido.strip())
    if n_chars == 0:
        return "vacia"
    if n_chars < 200:
        return "parcial"
    return "completa"


class DocxReader:
    """Lector de archivos .docx que produce un `Documento` poblado.

    Uso:
        reader = DocxReader()
        documento = reader.leer(Path("ruta/al/archivo.docx"))
    """

    def leer(self, ruta: Path, user_id: str = "default") -> Documento:
        """Parsea el .docx y devuelve un `Documento` con secciones detectadas."""
        if not ruta.exists():
            raise FileNotFoundError(f"No existe el archivo: {ruta}")

        docx = DocxDocument(str(ruta))
        documento = Documento(
            user_id=user_id,
            archivo_origen=str(ruta),
            secciones=construir_secciones_vacias(),
        )

        # Recorrer párrafos acumulando contenido por sección detectada.
        seccion_actual: Seccion | None = None
        buffer: list[str] = []

        def flush_buffer() -> None:
            """Vuelca el buffer acumulado al contenido de la sección actual."""
            if seccion_actual is None:
                return
            contenido = "\n".join(b for b in buffer if b.strip())
            if not contenido.strip():
                return
            existente = seccion_actual.contenido or ""
            seccion_actual.contenido = (
                f"{existente}\n{contenido}".strip() if existente else contenido
            )
            seccion_actual.completitud = _evaluar_completitud(  # type: ignore[assignment]
                seccion_actual.contenido or ""
            )

        for parrafo in docx.paragraphs:
            texto = parrafo.text.strip()
            estilo = parrafo.style.name if parrafo.style else None

            if _es_heading(estilo):
                # Antes de cambiar de sección, vuelca lo acumulado a la actual.
                flush_buffer()
                buffer = []
                cat = _coincide_con_catalogo(texto)
                seccion_actual = documento.seccion_por_id(cat.id) if cat is not None else None
            else:
                if texto:
                    buffer.append(texto)

        # Flush final
        flush_buffer()

        # Capturar contenido de tablas y agregarlo a la última sección detectada
        # (heurística simple — para MVP es suficiente).
        for tabla in docx.tables:
            filas_texto: list[str] = []
            for fila in tabla.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if celdas:
                    filas_texto.append(" | ".join(celdas))
            if filas_texto and seccion_actual is not None:
                tabla_str = "\n".join(filas_texto)
                seccion_actual.contenido = (
                    f"{seccion_actual.contenido}\n\n[Tabla extraída]\n{tabla_str}"
                    if seccion_actual.contenido
                    else f"[Tabla extraída]\n{tabla_str}"
                )

        # Re-evaluar completitud de todas las secciones tras absorber tablas.
        for s in documento.secciones:
            s.completitud = _evaluar_completitud(s.contenido or "")  # type: ignore[assignment]

        # Audit trail: documento importado
        documento.registrar_evento(
            EventoAuditoria(
                timestamp=datetime.now(UTC),
                actor=user_id,
                tipo="documento_importado",
                descripcion=f"Importado desde {ruta.name}",
                metadata={"archivo": ruta.name},
            )
        )

        return documento

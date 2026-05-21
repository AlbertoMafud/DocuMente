"""Lector de DOCX secundario — extrae texto + opcionalmente imágenes embebidas.

A diferencia del `DocxReader` principal de `src/docs/reader.py` que parsea
estructura del template NYL, este reader solo extrae párrafos y celdas de
tablas como texto, sin interpretar headings ni mapear a secciones.

Útil cuando el usuario adjunta un `.docx` como **fuente de contexto**
adicional (instructivos, procedimientos viejos, notas técnicas) en lugar
de como ancla estructural del documento.

Si se pasa un `VisionDescriber`, también extrae imágenes embebidas
(las que viven en `word/media/*.png|jpg|...` dentro del zip del .docx),
las describe con Claude Vision e intercala las descripciones al final
del texto extraído.
"""

from __future__ import annotations

import logging
from typing import IO

from docx import Document as DocxDocument

from src.llm.vision_describer import VisionDescriber

logger = logging.getLogger(__name__)

# Imágenes que `python-docx` reporta en `inline_shapes` pueden estar en
# distintos formatos; el `image.content_type` da el MIME.


def leer_docx_texto(
    archivo: IO[bytes],
    *,
    vision_describer: VisionDescriber | None = None,
) -> str:
    """Devuelve párrafos + tablas como texto plano, separados por saltos de línea.

    Si `vision_describer` se pasa, también describe imágenes embebidas y
    añade las descripciones al final, una por línea, como:
    `[Imagen N: <descripción>]`. No se intenta preservar el orden exacto
    relativo al texto (eso requeriría parsear el XML) — el LLM downstream
    típicamente puede correlacionar.
    """
    archivo.seek(0)
    doc = DocxDocument(archivo)

    fragmentos: list[str] = []

    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        if texto:
            fragmentos.append(texto)

    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                fragmentos.append(" | ".join(celdas))

    # Imágenes embebidas (inline shapes) — opcional
    if vision_describer is not None:
        descripciones = _describir_imagenes_docx(doc, vision_describer)
        if descripciones:
            fragmentos.append("")  # separador visual
            fragmentos.extend(descripciones)

    return "\n".join(fragmentos)


def _describir_imagenes_docx(doc: object, vision_describer: VisionDescriber) -> list[str]:
    """Itera las imágenes embebidas del .docx y devuelve descripciones.

    Recorre `doc.part.related_parts` buscando ImageParts. Esto incluye
    imágenes embebidas via inline_shapes (la forma común) y también las
    embebidas como floating shapes.
    """
    descripciones: list[str] = []
    contador = 0

    # `doc.part.related_parts` es un dict id→Part de python-docx
    related = getattr(getattr(doc, "part", None), "related_parts", {}) or {}
    for _rel_id, part in related.items():
        content_type = getattr(part, "content_type", "")
        if not content_type.startswith("image/"):
            continue
        try:
            imagen_bytes = part.blob  # bytes del binario embebido
        except Exception as exc:
            logger.debug("No se pudo leer blob de imagen %s: %s", _rel_id, exc)
            continue
        if not imagen_bytes:
            continue

        contador += 1
        desc = vision_describer.describir(imagen_bytes, media_type=content_type)
        descripciones.append(f"[Imagen {contador}: {desc.descripcion}]")

    return descripciones

"""Use case: ExtraerEstructuraTemplate (Template Studio, paso 2).

Recorre un .docx "hecho para humanos" con python-docx y detecta el esqueleto:
encabezados (por estilo Heading/Título o numeración tipo '4.1 '), el texto
agregado bajo cada encabezado y las tablas. El resultado se muestra al admin
para verificación visual ANTES de gastar LLM (spec §4 paso 2).

Sin LLM: esto es determinístico a propósito.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import IO

from docx import Document as DocxDocument

_NUMERACION_RE = re.compile(r"^(\d+(?:\.\d+)*)[.)]?\s+(.{2,})$")
_MAX_CHARS_TEXTO_POR_SECCION = 4000


@dataclass
class SeccionCandidata:
    """Un encabezado detectado + su contenido agregado."""

    titulo: str
    numero: str  # "" si el encabezado no traía numeración
    nivel: int  # 1 = principal
    texto: str = ""
    n_tablas: int = 0


@dataclass
class EstructuraExtraida:
    """Esqueleto del template detectado en el .docx."""

    secciones: list[SeccionCandidata] = field(default_factory=list)
    preambulo: str = ""
    """Texto antes del primer encabezado (portada, propósito general)."""
    n_tablas_total: int = 0
    advertencias: list[str] = field(default_factory=list)


def _nivel_de_estilo(nombre_estilo: str) -> int | None:
    """Devuelve el nivel si el estilo es un heading de Word, None si no."""
    m = re.match(r"^(heading|título|titulo)\s*(\d+)$", nombre_estilo.strip().lower())
    if m:
        return int(m.group(2))
    return None


class ExtraerEstructuraTemplate:
    """Extrae el esqueleto de un .docx de template."""

    def ejecutar(self, archivo: IO[bytes]) -> EstructuraExtraida:
        doc = DocxDocument(archivo)
        resultado = EstructuraExtraida()
        actual: SeccionCandidata | None = None
        buffer_preambulo: list[str] = []

        for parrafo in doc.paragraphs:
            texto = parrafo.text.strip()
            if not texto:
                continue

            estilo = parrafo.style.name if parrafo.style is not None else ""
            nivel_estilo = _nivel_de_estilo(estilo)
            m_num = _NUMERACION_RE.match(texto)

            es_heading = nivel_estilo is not None or (
                # Numeración al inicio + línea corta = encabezado probable.
                m_num is not None and len(texto) <= 120
            )

            if es_heading:
                if m_num:
                    numero, titulo = m_num.group(1), m_num.group(2).strip()
                    nivel = nivel_estilo or numero.count(".") + 1
                else:
                    numero, titulo = "", texto
                    nivel = nivel_estilo or 1
                actual = SeccionCandidata(titulo=titulo, numero=numero, nivel=nivel)
                resultado.secciones.append(actual)
            elif actual is not None:
                if len(actual.texto) < _MAX_CHARS_TEXTO_POR_SECCION:
                    actual.texto = f"{actual.texto}\n{texto}".strip()
            else:
                buffer_preambulo.append(texto)

        # Tablas: python-docx las expone a nivel documento; sin posición exacta
        # respecto a los párrafos (limitación conocida). Contamos el total y lo
        # anotamos — la propuesta LLM y la curación deciden a qué sección van.
        resultado.n_tablas_total = len(doc.tables)
        if doc.tables and resultado.secciones:
            resultado.advertencias.append(
                f"El documento tiene {len(doc.tables)} tabla(s); revisa en la curación "
                "qué secciones deben ser tipo 'tabla'."
            )

        resultado.preambulo = "\n".join(buffer_preambulo)[:_MAX_CHARS_TEXTO_POR_SECCION]

        if not resultado.secciones:
            resultado.advertencias.append(
                "No se detectaron encabezados (estilos Heading/Título ni numeración). "
                "El LLM propondrá estructura solo desde el texto — revisa con más cuidado."
            )
        return resultado

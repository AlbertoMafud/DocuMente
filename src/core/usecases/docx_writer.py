"""DocxWriter — render del Documento contra plantilla maestra `.docx`.

Toma un `Documento` Pydantic y la plantilla maestra editada por el usuario
(con placeholders Jinja en `src/docs/templates/model_development_smnyl_final.docx`).
Construye el contexto de variables, renderiza con `docxtpl` y devuelve los
bytes del DOCX listo para descargar.

Política de calidad estética: este módulo NO genera estilos en código. Toda
la marca SMNYL (paleta, tipografías, encabezados, layout) vive en la plantilla
Word. Aquí solo rellenamos placeholders.

Para las 4 secciones tabulares del template (5.1, 5.2, 5.5, 6.5) usa
`TableExtractor` con Haiku para convertir el texto narrativo en estructura
JSON. Si no se inyecta extractor, los loops salen vacíos.
"""

from __future__ import annotations

import contextlib
import re
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import Final

import docx as docx_lib
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import Table as DocxTable
from docxtpl import DocxTemplate
from docxtpl.subdoc import Subdoc

from src.core.models import Documento
from src.core.models.documento import EstadoDocumento
from src.core.usecases.markdown_blocks import (
    BloqueProsa,
    BloqueTabla,
    font_size_para_tabla,
    separar_bloques,
)
from src.core.usecases.markdown_cleanup import limpiar_markdown
from src.core.usecases.richtext_render import ParrafoSpec, parsear_parrafos
from src.core.usecases.strings_localizados import Idioma, t
from src.core.usecases.table_extractor import TableExtractor, TableSchema

# Mapeo seccion_id (catálogo NYL) → placeholder simple (definido por el usuario en plantilla)
_MAPA_SECCION_PLACEHOLDER: Final[dict[str, str]] = {
    "1.3.problem_statement": "problem_statement",
    "2.1.model_uses": "model_uses",
    "2.2.model_scope": "model_scope",
    "2.3.business_impact": "business_impact",
    "3.1.ancillary": "ancillary_docs",
    "3.2.additional": "additional_docs",
    "4.1.diagram": "schematic_description",
    "4.2.theory": "theory_and_logic",
    "4.3.risk_drivers": "key_risk_drivers",
    "4.4.assumptions": "key_assumptions",
    "5.3.1.aggregations": "data_aggregations",
    "5.3.2.segmentations": "segmentations",
    "5.3.3.averages_proxies": "averages_proxies",
    "5.4.data_limitations": "data_limitations",
    "6.1.specification": "specification_process",
    "6.2.approach": "approach_used",
    "6.3.dev_testing": "development_testing",
    "6.4.limitations": "limitations_revealed",
    "7.1.platform": "platform",
    "7.2.runs": "model_runs",
    "7.3.perf_testing": "performance_testing",
    "7.4.prod_limitations": "production_limitations",
    "8.governance": "governance",
    "9.monitoring": "ongoing_monitoring",
}

# Secciones que en el template NYL son tablas estructuradas, no texto plano.
# Estas requieren TableExtractor para llenar el loop correspondiente.
SCHEMA_RAW_DATA_SOURCES: Final = TableSchema(
    nombre="raw_data_sources",
    campos=["data_source", "data_type", "location", "method", "team"],
    descripcion_para_llm=(
        "Fuentes de datos crudos consumidas por el modelo. Cada fila: "
        "nombre de la fuente, tipo de dato (tabla, archivo, API…), ubicación "
        "o ruta, método de carga y equipo responsable."
    ),
)
SCHEMA_UPSTREAM_MODELS: Final = TableSchema(
    nombre="upstream_models",
    campos=["num", "name", "key_contact", "inventory_id"],
    descripcion_para_llm=(
        "Modelos upstream o supuestos determinados por la compañía que el "
        "modelo consume. Cada fila: número secuencial, nombre, contacto clave "
        "y model ID en inventario MRM."
    ),
)
SCHEMA_INPUT_CHANGES: Final = TableSchema(
    nombre="input_changes",
    campos=["date", "decision", "change_description"],
    descripcion_para_llm=(
        "Cambios o decisiones tomadas sobre los inputs del modelo. Cada fila: "
        "fecha del cambio, decisión tomada y descripción del cambio."
    ),
)
SCHEMA_PROCESS_CHANGES: Final = TableSchema(
    nombre="process_changes",
    campos=["date", "decision", "change_description"],
    descripcion_para_llm=(
        "Cambios al proceso de construcción del modelo. Cada fila: fecha, "
        "decisión tomada, descripción del cambio."
    ),
)

# Tabla de extracciones: (seccion_id_origen, schema, nombre del placeholder de loop)
_TABLAS_TABULARES: Final[list[tuple[str, TableSchema]]] = [
    ("5.1.raw_data", SCHEMA_RAW_DATA_SOURCES),
    ("5.2.upstream", SCHEMA_UPSTREAM_MODELS),
    ("5.5.input_changes", SCHEMA_INPUT_CHANGES),
    ("6.5.process_changes", SCHEMA_PROCESS_CHANGES),
]

_ETIQUETA_ESTADO_HUMANA: Final[dict[EstadoDocumento, str]] = {
    "draft": "Borrador",
    "in_review": "En revisión",
    "approved": "Aprobado",
    "published": "Publicado",
    "retired": "Retirado",
}


class DocxWriter:
    """Render del Documento al .docx final usando docxtpl + plantilla SMNYL."""

    def __init__(self, table_extractor: TableExtractor | None = None) -> None:
        self.table_extractor = table_extractor

    def generar(
        self,
        documento: Documento,
        template_path: Path,
        *,
        idioma: Idioma = "es",
    ) -> bytes:
        """Renderiza la plantilla con los datos del documento y devuelve bytes del .docx.

        `idioma` controla las cadenas generadas por el writer (marcadores de
        sección omitida, "Pendiente", etc.). El contenido propiamente dicho
        debe traducirse antes vía `TraductorDocumento`.

        Apéndices: se acumulan al final del documento en una sección dedicada
        "Apéndices" / "Appendix", numerados A.1, A.2…. Las referencias
        `(ver Apéndice: <titulo>)` o `(see Appendix: <titulo>)` en el body
        se reemplazan automáticamente por `(ver Apéndice A.N)` /
        `(see Appendix A.N)` según el índice asignado.
        """
        tpl = DocxTemplate(template_path)
        contexto = self._construir_contexto(documento, tpl, idioma=idioma)
        tpl.render(contexto)
        buffer = BytesIO()
        tpl.save(buffer)

        # Si hay apéndices, los agregamos al final reabriendo el .docx con
        # python-docx (los Subdocs de docxtpl no soportan heading styles ni
        # secciones independientes del template).
        if documento.apendices:
            buffer.seek(0)
            doc_final = docx_lib.Document(buffer)
            _agregar_seccion_apendices(doc_final, documento, idioma)
            buffer_final = BytesIO()
            doc_final.save(buffer_final)
            return buffer_final.getvalue()

        return buffer.getvalue()

    def _construir_contexto(
        self,
        documento: Documento,
        tpl: DocxTemplate,
        *,
        idioma: Idioma = "es",
    ) -> dict[str, object]:
        meta = documento.metadata_modelo
        ahora = datetime.now(UTC).astimezone()
        contexto: dict[str, object] = {
            # Metadata simple
            "nombre_modelo": meta.nombre_modelo or "—",
            "model_id": meta.model_id or "—",
            "model_class": meta.model_class or "—",
            "profit_center": meta.profit_center or "—",
            "bu_executive": meta.fae or "—",
            "model_owner": meta.model_owner or "—",
            "model_developers": ", ".join(meta.model_developers) if meta.model_developers else "—",
            "current_version": meta.current_version or "1.0",
            "version_actual": meta.current_version or "1.0",
            "implementation_platform": meta.implementation_platform or "—",
            "financial_impact": meta.financial_impact or "—",
            "model_status": meta.model_status or "—",
            "target_production_date": meta.target_production_date or "—",
            # Portada / encabezado
            "autor": meta.model_owner or "—",
            "fecha_documentacion": ahora.strftime("%Y-%m-%d"),
            "estado_documento": _ETIQUETA_ESTADO_HUMANA[documento.estado],
        }

        # Índice de apéndices por título → "A.N" para reemplazar cross-refs
        # en el body antes de renderizar.
        apendice_label_por_titulo = _indexar_apendices(documento, idioma)

        # Mapeo de placeholders de sección — cada uno como Subdoc con formato real.
        # Los apéndices YA NO se inyectan en el Subdoc de la sección; se agregan
        # como sección dedicada al final del documento (ver `generar`).
        for seccion_id, placeholder in _MAPA_SECCION_PLACEHOLDER.items():
            seccion = documento.seccion_por_id(seccion_id)
            contexto[placeholder] = _renderizar_seccion(
                tpl,
                seccion,
                idioma=idioma,
                apendice_label_por_titulo=apendice_label_por_titulo,
            )

        # Loop version_history desde audit_trail
        contexto["version_history"] = _construir_version_history(documento)

        # Loops tabulares (vía TableExtractor o vacíos si no hay extractor)
        for seccion_id, schema in _TABLAS_TABULARES:
            if self.table_extractor is None:
                contexto[schema.nombre] = []
            else:
                contexto[schema.nombre] = self.table_extractor.extraer(
                    documento, seccion_id, schema
                )

        return contexto


def _renderizar_seccion(
    tpl: DocxTemplate,
    seccion: object,
    *,
    idioma: Idioma = "es",
    apendice_label_por_titulo: dict[str, str] | None = None,
) -> Subdoc:
    """Devuelve un Subdoc con el contenido de la sección.

    Procesamiento:
    1. Reemplazar referencias `(ver Apéndice: <titulo>)` por `(ver Apéndice A.N)`
       (idem para inglés) si `apendice_label_por_titulo` mapea el título.
    2. `limpiar_markdown` quita tablas pipe-separated, separadores y hashes
       (preservando `**bold**` y `*italic*` para conversión a formato real).
    3. `parsear_parrafos` divide en bloques con runs.
    4. Cada `ParrafoSpec` se vuelve un párrafo del Subdoc con runs reales.
    5. Subtítulos (líneas solas con `**xxx**`) alineados a la izquierda.
    6. Bullets reciben prefijo `•`.

    Los apéndices YA NO se rendereaban aquí — ahora viven en una sección
    dedicada al final del documento (ver `_agregar_seccion_apendices`).
    """
    sub = tpl.new_subdoc()
    if seccion is None:
        sub.add_paragraph(t("seccion_no_catalogo", idioma))
        return sub

    completitud = getattr(seccion, "completitud", None)
    contenido = getattr(seccion, "contenido", None) or ""
    motivo = getattr(seccion, "motivo_omision", None) or ""

    if completitud == "omitida":
        razon = motivo or t("seccion_sin_motivo", idioma)
        sub.add_paragraph(f"{t('seccion_omitida_prefijo', idioma)}{razon}")
        return sub
    if not contenido.strip():
        sub.add_paragraph(t("pendiente_sin_contenido", idioma))
        return sub

    if apendice_label_por_titulo:
        contenido = _reemplazar_referencias_apendices(contenido, apendice_label_por_titulo, idioma)

    texto = limpiar_markdown(contenido, conservar_enfasis=True)
    parrafos = parsear_parrafos(texto)
    if not parrafos:
        sub.add_paragraph(texto)
    else:
        for p_spec in parrafos:
            _agregar_parrafo(sub, p_spec)

    return sub


def _indexar_apendices(documento: Documento, idioma: Idioma) -> dict[str, str]:
    """Mapea cada título de apéndice a su etiqueta de cross-ref (`A.1`, `A.2`…).

    Numeración por orden de aparición en `documento.apendices` (orden de
    adjunción del usuario). Si dos apéndices comparten título, gana el primero.
    Devuelve `{}` si no hay apéndices.

    El idioma no afecta la numeración pero se acepta como parámetro para
    futura extensibilidad (ej. numeración localizada).
    """
    del idioma  # numeración A.N es universal; idioma reservado para extensión
    mapa: dict[str, str] = {}
    for i, ap in enumerate(documento.apendices, start=1):
        titulo = (getattr(ap, "titulo", "") or "").strip()
        if titulo and titulo not in mapa:
            mapa[titulo] = f"A.{i}"
    return mapa


_REGEX_REF_APENDICE = re.compile(
    r"\((ver Apéndice|see Appendix):\s*([^)]+?)\)",
    re.IGNORECASE,
)


def _reemplazar_referencias_apendices(
    contenido: str,
    label_por_titulo: dict[str, str],
    idioma: Idioma,
) -> str:
    """Reemplaza `(ver Apéndice: <titulo>)` por `(ver Apéndice A.N)`.

    En inglés, también acepta `(see Appendix: <titulo>)` y produce
    `(see Appendix A.N)`. Si el título no matchea ningún apéndice, deja
    la referencia tal cual (no rompe).
    """
    if not label_por_titulo:
        return contenido

    etiqueta = t("ver_apendice", idioma)  # "ver Apéndice" o "see Appendix"

    def _sustituir(match: re.Match[str]) -> str:
        titulo_referenciado = match.group(2).strip()
        label = label_por_titulo.get(titulo_referenciado)
        if label is None:
            # Búsqueda case-insensitive como fallback
            for titulo_real, lbl in label_por_titulo.items():
                if titulo_real.lower() == titulo_referenciado.lower():
                    label = lbl
                    break
        if label is None:
            return match.group(0)  # sin match, dejar tal cual
        return f"({etiqueta} {label})"

    return _REGEX_REF_APENDICE.sub(_sustituir, contenido)


def _agregar_seccion_apendices(
    doc: DocxDocument,
    documento: Documento,
    idioma: Idioma,
) -> None:
    """Agrega al final del documento la sección "Apéndices" / "Appendix".

    Cada apéndice:
    - Heading nivel 2: "A.N: <titulo>".
    - Contenido bloque por bloque: tablas markdown → tablas nativas con
      bordes y font adaptable; prosa → párrafos con runs (bold/italic).
    """
    if not documento.apendices:
        return

    # Encabezado de sección.
    doc.add_heading(t("apendices_plural", idioma), level=1)

    for i, ap in enumerate(documento.apendices, start=1):
        titulo = (getattr(ap, "titulo", "") or "").strip() or t("apendice_singular", idioma)
        contenido_md = getattr(ap, "contenido_md", "") or ""

        # Heading "A.N: <titulo>"
        doc.add_heading(f"A.{i}: {titulo}", level=2)

        if not contenido_md.strip():
            continue

        for bloque in separar_bloques(contenido_md):
            if isinstance(bloque, BloqueTabla):
                _agregar_tabla_documento(doc, bloque)
            elif isinstance(bloque, BloqueProsa):
                texto_limpio = limpiar_markdown(bloque.texto, conservar_enfasis=True)
                for p_spec in parsear_parrafos(texto_limpio):
                    _agregar_parrafo_documento(doc, p_spec)


def _agregar_tabla_documento(doc: DocxDocument, bloque: BloqueTabla) -> None:
    """Versión de `_agregar_tabla_word` que opera sobre un Document, no Subdoc."""
    if not bloque.headers:
        return
    n_columnas = len(bloque.headers)
    n_filas_datos = len(bloque.rows)
    font_pt = font_size_para_tabla(n_filas=n_filas_datos, n_columnas=n_columnas)

    table: DocxTable = doc.add_table(rows=1 + n_filas_datos, cols=n_columnas)
    with contextlib.suppress(KeyError):
        table.style = "Table Grid"

    for j, header in enumerate(bloque.headers):
        cell = table.cell(0, j)
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(font_pt)

    for i, fila in enumerate(bloque.rows, start=1):
        for j in range(n_columnas):
            valor = fila[j] if j < len(fila) else ""
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(valor)
            run.font.size = Pt(font_pt)


def _agregar_parrafo_documento(doc: DocxDocument, p_spec: ParrafoSpec) -> None:
    """Versión de `_agregar_parrafo` que opera sobre un Document, no Subdoc."""
    prefijo = "•  " if p_spec.es_bullet else ""

    parrafo = doc.add_paragraph()
    if p_spec.es_subtitulo:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if prefijo:
        parrafo.add_run(prefijo)

    for run_spec in p_spec.runs:
        run = parrafo.add_run(run_spec.text)
        if run_spec.bold:
            run.bold = True
        if run_spec.italic:
            run.italic = True


def _agregar_tabla_word(sub: Subdoc, bloque: BloqueTabla) -> None:
    """Inserta una tabla nativa de Word con bordes y font size adaptable."""
    if not bloque.headers:
        return
    n_columnas = len(bloque.headers)
    n_filas_datos = len(bloque.rows)
    font_pt = font_size_para_tabla(n_filas=n_filas_datos, n_columnas=n_columnas)

    table = sub.add_table(rows=1 + n_filas_datos, cols=n_columnas)
    # Si la plantilla no tiene 'Table Grid', el estilo default queda y no hay bordes.
    with contextlib.suppress(KeyError):
        table.style = "Table Grid"

    # Header
    for j, header in enumerate(bloque.headers):
        cell = table.cell(0, j)
        cell.text = ""  # limpiar default
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(font_pt)

    # Filas de datos
    for i, fila in enumerate(bloque.rows, start=1):
        for j in range(n_columnas):
            valor = fila[j] if j < len(fila) else ""
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(valor)
            run.font.size = Pt(font_pt)


def _agregar_parrafo(sub: Subdoc, p_spec: ParrafoSpec) -> None:
    """Agrega un párrafo al Subdoc respetando los runs y propiedades."""
    # Prefijo visual; en iteración posterior podría usarse list style real de Word.
    prefijo = "•  " if p_spec.es_bullet else ""

    parrafo = sub.add_paragraph()
    if p_spec.es_subtitulo:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if prefijo:
        parrafo.add_run(prefijo)

    for run_spec in p_spec.runs:
        run = parrafo.add_run(run_spec.text)
        if run_spec.bold:
            run.bold = True
        if run_spec.italic:
            run.italic = True


def _construir_version_history(documento: Documento) -> list[dict[str, str]]:
    """Construye la lista para el loop {%tr for v in version_history %}.

    Estrategia MVP: cada transición de estado en el audit_trail se vuelve una
    fila. La 'versión' se incrementa secuencialmente. Si no hay transiciones,
    se devuelve una sola fila con la versión inicial 1.0 a partir del primer
    evento de creación/importación.
    """
    transiciones = [e for e in documento.audit_trail if e.tipo == "transicion_estado"]
    if not transiciones:
        primer_evento = next(iter(documento.audit_trail), None)
        if primer_evento is None:
            return []
        return [
            {
                "version_no": "1.0",
                "date_changed": primer_evento.timestamp.astimezone().strftime("%Y-%m-%d"),
                "updated_by": primer_evento.actor,
                "approved_by": "—",
                "description": primer_evento.descripcion,
            }
        ]

    filas: list[dict[str, str]] = []
    for i, evento in enumerate(transiciones, start=1):
        filas.append(
            {
                "version_no": f"{i}.0",
                "date_changed": evento.timestamp.astimezone().strftime("%Y-%m-%d"),
                "updated_by": evento.actor,
                "approved_by": "—",
                "description": evento.descripcion,
            }
        )
    return filas

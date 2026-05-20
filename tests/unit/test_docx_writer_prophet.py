from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from src.core.models import Documento
from src.core.models.documento import MetadataModelo
from src.core.template_catalog_prophet import construir_secciones_vacias_prophet
from src.core.usecases.docx_writer_prophet import DocxWriterProphet


def _doc_prophet_con_runs(nombre: str = "Modelo VNB") -> Documento:
    secciones = construir_secciones_vacias_prophet()
    runs_data = json.dumps(
        {
            "filas": [
                {
                    "numero": "33",
                    "detalle": "IL UDI",
                    "es_alm": "No",
                    "tiempo_ejecucion": "45 min",
                    "corrida_precedente": "",
                    "outputs_principales": "VNB",
                    "responsable": "Carmona",
                },
            ],
            "advertencias": [],
        }
    )
    for s in secciones:
        if s.id == "corridas_runs":
            s.contenido = runs_data
            s.completitud = "completa"
    return Documento(
        tipo="prophet",
        metadata_modelo=MetadataModelo(nombre_modelo=nombre),
        secciones=secciones,
    )


def _template_minimo(tmp_path: Path) -> Path:
    """Crea un .docx mínimo con un placeholder para tests."""
    doc = DocxDocument()
    doc.add_paragraph("{{ nombre_modelo }}")
    doc.add_paragraph("{{ area }}")
    path = tmp_path / "prophet_test_template.docx"
    doc.save(str(path))
    return path


def test_construir_contexto_incluye_nombre_modelo(tmp_path: Path) -> None:
    writer = DocxWriterProphet(template_path=_template_minimo(tmp_path))
    doc = _doc_prophet_con_runs("Modelo VNB")
    ctx = writer._construir_contexto(doc)
    assert ctx["nombre_modelo"] == "Modelo VNB"


def test_construir_contexto_runs_como_lista(tmp_path: Path) -> None:
    writer = DocxWriterProphet(template_path=_template_minimo(tmp_path))
    doc = _doc_prophet_con_runs()
    ctx = writer._construir_contexto(doc)
    assert isinstance(ctx["runs"], list)
    assert len(ctx["runs"]) == 1
    assert ctx["runs"][0]["numero"] == "33"


def test_construir_contexto_seccion_vacia_devuelve_lista_vacia(tmp_path: Path) -> None:
    writer = DocxWriterProphet(template_path=_template_minimo(tmp_path))
    doc = Documento(
        tipo="prophet",
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=construir_secciones_vacias_prophet(),
    )
    ctx = writer._construir_contexto(doc)
    assert ctx["runs"] == []
    assert ctx["variables"] == []


def test_render_devuelve_bytes(tmp_path: Path) -> None:
    writer = DocxWriterProphet(template_path=_template_minimo(tmp_path))
    doc = _doc_prophet_con_runs()
    result = writer.render(doc)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_template_path_no_existe_lanza_file_not_found(tmp_path: Path) -> None:
    writer = DocxWriterProphet(template_path=tmp_path / "no_existe.docx")
    doc = _doc_prophet_con_runs()
    with pytest.raises(FileNotFoundError):
        writer.render(doc)

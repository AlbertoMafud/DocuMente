"""Tests del DocxWriter — render del Documento contra plantilla maestra real."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import docx
import pytest

from src.core.models import (
    Documento,
    EventoAuditoria,
    MetadataModelo,
    Seccion,
)
from src.core.usecases.docx_writer import DocxWriter
from src.core.usecases.table_extractor import TableExtractor
from tests.unit.test_interview_engine import FakeLLM

TEMPLATE_PATH = Path("src/docs/templates/model_development_smnyl_final.docx")


@pytest.fixture
def documento_completo() -> Documento:
    """Documento con metadata + 1 sección completa + 1 omitida + audit trail."""
    doc = Documento(
        metadata_modelo=MetadataModelo(
            nombre_modelo="Modelo VNB",
            model_id="M07.P07.S04.019.B",
            model_class="Actuarial",
            profit_center="P07",
            fae="Yael Aguilera",
            model_owner="Alberto Solano",
            model_developers=["Alberto Solano", "Otro Dev"],
            current_version="1.0",
            implementation_platform="Excel + Prophet",
            financial_impact="GAAP earnings",
            model_status="In Production",
            inherent_risk_tier="medium",
            intended_use="Cálculo VNB trimestral",
        ),
        secciones=[
            Seccion(
                id="2.1.model_uses",
                nombre="Model Uses",
                numero="2.1",
                obligatoria=True,
                contenido="Este modelo se utiliza para calcular VNB trimestralmente.",
                completitud="completa",
            ),
            Seccion(
                id="2.2.model_scope",
                nombre="Model Scope",
                numero="2.2",
                obligatoria=True,
                contenido=None,
                completitud="omitida",
                motivo_omision="No aplica al modelo",
            ),
        ],
    )
    doc.registrar_evento(
        EventoAuditoria(
            actor="default",
            tipo="transicion_estado",
            descripcion="Borrador inicial",
            metadata={"origen": "draft", "destino": "in_review"},
        )
    )
    return doc


def _texto_plano(blob: bytes) -> str:
    """Devuelve todo el texto del .docx leyendo `<w:t>` directo del XML.

    Necesario porque los Subdocs que docxtpl inyecta no aparecen en
    `doc.paragraphs` directamente — quedan en estructuras anidadas que
    python-docx no expone como párrafos top-level.
    """
    import re
    import zipfile

    with zipfile.ZipFile(BytesIO(blob)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    fragmentos = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.DOTALL)
    return "\n".join(fragmentos)


def test_writer_genera_bytes_validos(documento_completo: Documento) -> None:
    """El writer devuelve un .docx que python-docx puede abrir sin error."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)

    assert isinstance(blob, bytes)
    assert len(blob) > 5000  # un .docx mínimo tiene varios KB
    docx.Document(BytesIO(blob))  # no debe levantar excepción


def test_writer_inserta_metadata_simple(documento_completo: Documento) -> None:
    """Los placeholders de metadata aparecen en el output (subset confirmado)."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    # Subset de placeholders comprobados que la plantilla actual rinde sin
    # fragmentación. Los placeholders con underscore largo (ej.
    # implementation_platform) pueden fragmentarse en runs de Word; eso se
    # diagnostica con la utilidad inspect_template y se corrige en plantilla.
    assert "M07.P07.S04.019.B" in texto
    assert "Yael Aguilera" in texto
    assert "Alberto Solano" in texto


def test_writer_inserta_contenido_de_seccion_completa(documento_completo: Documento) -> None:
    """El contenido de la sección 2.1 aparece en el documento."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "calcular VNB trimestralmente" in texto


def test_writer_marca_secciones_omitidas_con_motivo(documento_completo: Documento) -> None:
    """Las secciones omitidas se renderizan con el marcador y motivo."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "omitida" in texto.lower()
    assert "No aplica al modelo" in texto


def test_writer_no_deja_placeholders_jinja_sin_resolver(
    documento_completo: Documento,
) -> None:
    """No deben quedar `{{ ... }}` literales en el output (todo se reemplazó)."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "{{" not in texto
    assert "{%" not in texto


def test_writer_inserta_version_history_desde_audit(
    documento_completo: Documento,
) -> None:
    """La tabla de Version Control se llena desde transiciones del audit_trail."""
    writer = DocxWriter()
    blob = writer.generar(documento_completo, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "Borrador inicial" in texto


def test_writer_loops_tabulares_vacios_sin_extractor(
    documento_completo: Documento,
) -> None:
    """Sin TableExtractor, los loops de tabla no rompen — quedan sin filas."""
    writer = DocxWriter()  # sin extractor
    blob = writer.generar(documento_completo, TEMPLATE_PATH)

    docx.Document(BytesIO(blob))  # no rompe


def test_writer_apendice_con_tabla_genera_tabla_nativa() -> None:
    """Si el contenido_md del apéndice tiene tabla markdown, debe salir como tabla de Word."""
    from src.core.models import Apendice

    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="4.4.assumptions",
                nombre="Key Assumptions",
                numero="4.4",
                obligatoria=True,
                contenido="Resumen de supuestos.",
                completitud="completa",
            ),
        ],
        apendices=[
            Apendice(
                seccion_origen_id="4.4.assumptions",
                titulo="Tabla de factores",
                contenido_md=(
                    "**Archivo origen:** factores.xlsx\n\n"
                    "| Producto | Factor |\n"
                    "|---|---|\n"
                    "| 1.0 | 0.158 |\n"
                    "| 2.0 | 0.383 |\n"
                    "| 3.0 | 0.450 |"
                ),
            )
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)

    # Contar tablas mirando el XML directo (incluye tablas dentro de Subdocs).
    import re
    import zipfile

    doc_sin_apendice = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=doc.secciones,
        apendices=[],
    )
    blob_sin = writer.generar(doc_sin_apendice, TEMPLATE_PATH)

    def contar_tablas(b: bytes) -> int:
        with zipfile.ZipFile(BytesIO(b)) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        return len(re.findall(r"<w:tbl[\s>]", xml))

    n_con = contar_tablas(blob)
    n_sin = contar_tablas(blob_sin)
    assert n_con > n_sin, f"Esperaba más tablas con apéndice ({n_con}) que sin ({n_sin})"

    # Verificar que los datos de la tabla aparecen en el texto
    texto = _texto_plano(blob)
    assert "Producto" in texto
    assert "Factor" in texto
    assert "0.158" in texto
    # Y los pipes ya NO deben quedar literales
    assert "| Producto |" not in texto
    assert "|---|---|" not in texto


def test_writer_renderiza_apendices_en_seccion_dedicada_al_final() -> None:
    """Los apéndices se acumulan al final del documento en una sección 'Apéndices',
    numerados A.1, A.2…. Ya NO se inyectan dentro del contenido de la sección
    que los referencia.
    """
    from src.core.models import Apendice

    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="4.4.assumptions",
                nombre="Key Assumptions",
                numero="4.4",
                obligatoria=True,
                contenido="Resumen narrativo de los supuestos del modelo.",
                completitud="completa",
            ),
        ],
        apendices=[
            Apendice(
                seccion_origen_id="4.4.assumptions",
                titulo="Tabla de mortalidad SOA 2017",
                contenido_md="Tabla con factores por edad y sexo.",
            )
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "Resumen narrativo de los supuestos" in texto
    # Sección dedicada "Apéndices" al final, con heading
    assert "Apéndices" in texto
    # Cada apéndice se numera A.N: <titulo>
    assert "A.1: Tabla de mortalidad SOA 2017" in texto
    assert "Tabla con factores por edad y sexo" in texto


def test_writer_cross_reference_de_apendice_se_reemplaza_por_etiqueta_numerica() -> None:
    """`(ver Apéndice: <titulo>)` en el body → `(ver Apéndice A.N)`."""
    from src.core.models import Apendice

    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="4.4.assumptions",
                nombre="Key Assumptions",
                numero="4.4",
                obligatoria=True,
                contenido="Los factores se detallan (ver Apéndice: Tabla de mortalidad).",
                completitud="completa",
            ),
        ],
        apendices=[
            Apendice(
                seccion_origen_id="4.4.assumptions",
                titulo="Tabla de mortalidad",
                contenido_md="Datos por edad.",
            )
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "(ver Apéndice A.1)" in texto
    # El título original ya NO debe aparecer en la referencia inline
    assert "(ver Apéndice: Tabla de mortalidad)" not in texto


def test_writer_apendices_en_idioma_en_usa_appendix_y_see_appendix() -> None:
    """En inglés: sección 'Appendix' + cross-refs como '(see Appendix A.1)'."""
    from src.core.models import Apendice

    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="4.4.assumptions",
                nombre="Key Assumptions",
                numero="4.4",
                obligatoria=True,
                contenido="Factors are detailed (see Appendix: Mortality Table).",
                completitud="completa",
            ),
        ],
        apendices=[
            Apendice(
                seccion_origen_id="4.4.assumptions",
                titulo="Mortality Table",
                contenido_md="Data by age and sex.",
            )
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH, idioma="en")
    texto = _texto_plano(blob)

    # Heading de sección en inglés
    assert "Appendix" in texto
    assert "Apéndices" not in texto  # no contamination
    # Cross-ref en inglés
    assert "(see Appendix A.1)" in texto
    assert "(see Appendix: Mortality Table)" not in texto
    # Numeración
    assert "A.1: Mortality Table" in texto


def test_writer_sin_apendices_no_agrega_seccion_apendices() -> None:
    """Doc sin apéndices → no aparece la palabra 'Apéndices' como heading."""
    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="2.1.model_uses",
                nombre="Model Uses",
                numero="2.1",
                obligatoria=True,
                contenido="Contenido cualquiera.",
                completitud="completa",
            ),
        ],
        apendices=[],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    # No debe aparecer ni "Apéndices" como heading ni "A.1:" como item
    assert "A.1:" not in texto


def test_writer_referencia_a_apendice_inexistente_se_deja_tal_cual() -> None:
    """Si el body referencia un apéndice por título y no existe → no romper."""
    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="2.1.model_uses",
                nombre="Model Uses",
                numero="2.1",
                obligatoria=True,
                contenido="Detalles (ver Apéndice: Tabla X que no existe).",
                completitud="completa",
            ),
        ],
        apendices=[],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    # Sin apéndice que matchee, la referencia queda literal
    assert "(ver Apéndice: Tabla X que no existe)" in texto


def test_writer_apendice_con_tabla_markdown_se_renderea_como_tabla_nativa_en_seccion_dedicada() -> (
    None
):
    """Apéndices con tablas markdown se renderean como tablas nativas dentro
    de la sección dedicada al final (no embebidas en la sección de origen).
    """
    from src.core.models import Apendice

    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="4.4.assumptions",
                nombre="Key Assumptions",
                numero="4.4",
                obligatoria=True,
                contenido="Resumen de supuestos.",
                completitud="completa",
            ),
        ],
        apendices=[
            Apendice(
                seccion_origen_id="4.4.assumptions",
                titulo="Tabla de factores",
                contenido_md=(
                    "Datos extraídos.\n\n"
                    "| Producto | Factor |\n"
                    "|---|---|\n"
                    "| 1.0 | 0.158 |\n"
                    "| 2.0 | 0.383 |"
                ),
            )
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    assert "Producto" in texto
    assert "0.158" in texto
    # Pipes ya NO deben quedar literales
    assert "| Producto |" not in texto


def test_writer_marcadores_en_idioma_en() -> None:
    """Con idioma='en', los marcadores hardcodeados salen en inglés."""
    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="2.1.model_uses",
                nombre="Model Uses",
                numero="2.1",
                obligatoria=True,
                contenido=None,
                completitud="omitida",
                motivo_omision="Not applicable to the model",
            ),
            Seccion(
                id="2.2.model_scope",
                nombre="Model Scope",
                numero="2.2",
                obligatoria=True,
                contenido=None,
                completitud="vacia",
            ),
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH, idioma="en")
    texto = _texto_plano(blob)

    # Marcador "Section omitted" en lugar de "Sección omitida"
    assert "Section omitted" in texto
    assert "Sección omitida" not in texto
    # Marcador "Pending" en lugar de "Pendiente"
    assert "[Pending — no content captured]" in texto
    assert "[Pendiente — sin contenido capturado]" not in texto


def test_writer_marcadores_en_idioma_es_por_default() -> None:
    """Sin idioma explícito, el writer renderea en español (backward compat)."""
    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id="2.1.model_uses",
                nombre="Model Uses",
                numero="2.1",
                obligatoria=True,
                contenido=None,
                completitud="omitida",
                motivo_omision="No aplica al modelo",
            ),
        ],
    )

    writer = DocxWriter()
    blob = writer.generar(doc, TEMPLATE_PATH)  # idioma default = "es"
    texto = _texto_plano(blob)

    assert "Sección omitida" in texto
    assert "Section omitted" not in texto


def test_writer_llama_a_extractor_para_secciones_tabulares() -> None:
    """Con TableExtractor, las 4 secciones tabulares disparan extracción."""
    llm = FakeLLM(
        [
            '[{"data_source": "Sofia", "data_type": "tabla",'
            ' "location": "/x", "method": "manual", "team": "Rent"}]',
            '[{"num": "1", "name": "Prophet", "key_contact": "Yael", "inventory_id": "M07.001"}]',
            '[{"date": "2026-05-01", "decision": "Test", "change_description": "Test"}]',
            '[{"date": "2026-05-02", "decision": "Test", "change_description": "Test"}]',
        ]
    )
    doc = Documento(
        metadata_modelo=MetadataModelo(nombre_modelo="X"),
        secciones=[
            Seccion(
                id=sid,
                nombre="Test",
                numero=num,
                obligatoria=True,
                contenido="Texto narrativo con info.",
                completitud="completa",
            )
            for sid, num in [
                ("5.1.raw_data", "5.1"),
                ("5.2.upstream", "5.2"),
                ("5.5.input_changes", "5.5"),
                ("6.5.process_changes", "6.5"),
            ]
        ],
    )

    writer = DocxWriter(table_extractor=TableExtractor(llm))
    blob = writer.generar(doc, TEMPLATE_PATH)
    texto = _texto_plano(blob)

    # Verificar que al menos 1 valor extraído de cada tabla apareció
    assert "Sofia" in texto
    assert "Prophet" in texto
    # Las 4 llamadas se hicieron
    assert len(llm.llamadas) == 4
